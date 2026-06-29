"""
Word (DOCX) Extractor
=====================
Parses Word documents using python-docx.

Expected format (same as the legacy DocxParser):
    Question: <text>
    A. <option>
    B. <option>
    ...
    Answer: A
    Marks: 2
"""

from __future__ import annotations

import logging

from app.services.import_engine.extractors.base import BaseExtractor
from app.services.import_engine.models import (
    DocumentType,
    ExtractedQuestion,
    ExtractionResult,
)

logger = logging.getLogger(__name__)


class WordExtractor(BaseExtractor):
    extractor_name = "word"

    def extract(self, filepath: str) -> ExtractionResult:
        result = self._make_result(DocumentType.WORD)
        result.extractor_name = self.extractor_name

        try:
            import docx
        except ImportError:
            result.errors.append("python-docx is not installed.")
            return result

        try:
            doc = docx.Document(filepath)
            full_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
            result.metadata["full_text"] = full_text
            result.questions = self._parse_paragraphs(doc.paragraphs)
        except Exception as exc:
            logger.exception("WordExtractor failed: %s", exc)
            result.errors.append(f"Extraction error: {exc}")

        return result

    # ------------------------------------------------------------------

    def _parse_paragraphs(self, paragraphs) -> list[ExtractedQuestion]:
        questions: list[ExtractedQuestion] = []
        current: dict | None = None

        def finalise():
            if current and len(current["options"]) >= 2:
                q = ExtractedQuestion(
                    question_text=current["question"].strip(),
                    options=current["options"][:6],
                    marks=current["marks"],
                    raw_text=current["raw"],
                )
                q.correct_option_index = self._resolve_answer_index(
                    current.get("correct_raw", ""), q.options
                )
                q.confidence = self._question_confidence(q)
                questions.append(q)

        for p in paragraphs:
            line = p.text.strip()
            if not line:
                continue

            ll = line.lower()

            if ll.startswith("question"):
                finalise()
                q_text = line.split(":", 1)[1].strip() if ":" in line else line
                current = {
                    "question": q_text,
                    "options": [],
                    "correct_raw": None,
                    "marks": 1.0,
                    "raw": line,
                }

            elif current is not None:
                if len(line) >= 2 and line[0].upper() in "ABCDEF" and line[1] in ".) ":
                    current["options"].append(line[2:].strip())
                    current["raw"] += "\n" + line
                elif ll.startswith("answer"):
                    current["correct_raw"] = line.split(":", 1)[1].strip() if ":" in line else ""
                elif ll.startswith("marks"):
                    current["marks"] = self._parse_marks(line.split(":", 1)[1].strip() if ":" in line else None)
                elif not current["options"] and not current.get("correct_raw"):
                    current["question"] += "\n" + line

        finalise()
        return questions
